import io
from typing import BinaryIO
import pdfplumber
from docx import Document as DocxDocument
import logging

from app.core.exceptions import FileProcessingException, InvalidFileTypeException

logger = logging.getLogger(__name__)


class FileProcessor:
    """
    Extracts text content from various file formats.
    Supports: PDF, DOCX, TXT
    """
    
    SUPPORTED_TYPES = ["pdf", "docx", "txt"]
    
    def extract_text(self, file: BinaryIO, file_type: str) -> str:
        """
        Extract text from uploaded file.
        
        Args:
            file: File object
            file_type: Type of file (pdf, docx, txt)
        
        Returns:
            str: Extracted text content
        
        Raises:
            InvalidFileTypeException: If file type is not supported
            FileProcessingException: If text extraction fails
        """
        file_type = file_type.lower()
        
        logger.info(f"Starting text extraction for file type: {file_type}")
        
        if file_type not in self.SUPPORTED_TYPES:
            logger.error(f"Unsupported file type: {file_type}")
            raise InvalidFileTypeException(
                message=f"Unsupported file type: {file_type}",
                details={"file_type": file_type, "supported_types": self.SUPPORTED_TYPES}
            )
        
        # Extract based on file type
        if file_type == "pdf":
            return self._extract_from_pdf(file)
        elif file_type == "docx":
            return self._extract_from_docx(file)
        elif file_type == "txt":
            return self._extract_from_txt(file)
    
    def _extract_from_pdf(self, file: BinaryIO) -> str:
        """
        Extract text from PDF file using pdfplumber.
        
        Args:
            file: PDF file object
        
        Returns:
            str: Extracted text
        
        Raises:
            FileProcessingException: If PDF extraction fails
        """
        logger.info("Extracting text from PDF using pdfplumber library")
        
        try:
            # Open PDF with pdfplumber
            with pdfplumber.open(file) as pdf:
                num_pages = len(pdf.pages)
                logger.info(f"pdfplumber: PDF has {num_pages} pages")
                
                # Extract text from all pages
                text_parts = []
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
                        logger.debug(f"pdfplumber: Extracted text from page {i}/{num_pages} ({len(text)} chars)")
                
                # Combine all pages
                full_text = "\n\n".join(text_parts)
                logger.info(f"pdfplumber: Successfully extracted {len(full_text)} characters from PDF")
                
                if not full_text.strip():
                    logger.warning("pdfplumber: PDF extraction resulted in empty text")
                    raise FileProcessingException(
                        message="PDF contains no extractable text (pdfplumber)",
                        details={"pages": num_pages, "library": "pdfplumber"}
                    )
                
                return self._clean_text(full_text)
        
        except FileProcessingException:
            raise
        except Exception as e:
            logger.error(f"pdfplumber: Error extracting text from PDF: {str(e)}")
            raise FileProcessingException(
                message="Failed to extract text from PDF using pdfplumber",
                details={"error": str(e), "library": "pdfplumber"}
            )
    
    def _extract_from_docx(self, file: BinaryIO) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file: DOCX file object
        
        Returns:
            str: Extracted text
        
        Raises:
            FileProcessingException: If DOCX extraction fails
        """
        logger.debug("Extracting text from DOCX")
        
        try:
            # Load document
            doc = DocxDocument(file)
            num_paragraphs = len(doc.paragraphs)
            logger.info(f"DOCX has {num_paragraphs} paragraphs")
            
            # Extract text from all paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Combine all paragraphs
            full_text = "\n\n".join(text_parts)
            logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
            
            if not full_text.strip():
                logger.warning("DOCX extraction resulted in empty text")
                raise FileProcessingException(
                    message="DOCX contains no extractable text",
                    details={"paragraphs": num_paragraphs}
                )
            
            return self._clean_text(full_text)
        
        except FileProcessingException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {str(e)}")
            raise FileProcessingException(
                message="Failed to extract text from DOCX",
                details={"error": str(e)}
            )
    
    def _extract_from_txt(self, file: BinaryIO) -> str:
        """
        Extract text from TXT file.
        
        Args:
            file: TXT file object
        
        Returns:
            str: Extracted text
        
        Raises:
            FileProcessingException: If TXT extraction fails
        """
        logger.debug("Extracting text from TXT")
        
        try:
            # Read and decode text
            content = file.read()
            
            # Try UTF-8 first, fallback to latin-1
            try:
                text = content.decode('utf-8')
                logger.debug("Decoded TXT file as UTF-8")
            except UnicodeDecodeError:
                text = content.decode('latin-1')
                logger.debug("Decoded TXT file as latin-1 (UTF-8 failed)")
            
            logger.info(f"Successfully extracted {len(text)} characters from TXT")
            
            if not text.strip():
                logger.warning("TXT file is empty")
                raise FileProcessingException(
                    message="TXT file contains no text",
                    details={}
                )
            
            return self._clean_text(text)
        
        except FileProcessingException:
            raise
        except Exception as e:
            logger.error(f"Error extracting text from TXT: {str(e)}")
            raise FileProcessingException(
                message="Failed to extract text from TXT",
                details={"error": str(e)}
            )
    
    def _clean_text(self, text: str) -> str:
        """
        Clean extracted text.
        
        Args:
            text: Raw extracted text
        
        Returns:
            str: Cleaned text
        """
        logger.debug(f"Cleaning text ({len(text)} chars before cleaning)")
        
        # Remove excessive whitespace
        lines = [line.strip() for line in text.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join with single newline
        cleaned = '\n'.join(lines)
        
        logger.debug(f"Text cleaned ({len(cleaned)} chars after cleaning)")
        
        return cleaned
